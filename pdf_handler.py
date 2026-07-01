"""
PDF/DOCX 处理器 — PyMuPDF 提取 PDF，python-docx 提取 Word

PDF 处理：
  文字 PDF → PyMuPDF 毫秒级提取
  扫描件 PDF → 检测无文字层，提示安装 EasyOCR

生产环境建议使用 marker（https://github.com/VikParuchuri/marker）
获得更好的表格 / 公式 / 扫描件处理能力。
"""

import io
import logging

import fitz

logger = logging.getLogger(__name__)


# ── PDF 提取（PyMuPDF） ──

def extract_pdf(pdf_path: str) -> str:
    """从 PDF 提取文字（PyMuPDF）"""
    doc = fitz.open(pdf_path)
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()

    text = "\n\n--- 换页 ---\n\n".join(pages)
    logger.info("提取文字：%d 字符", len(text))
    return text


def is_scanned_pdf(pdf_path: str) -> bool:
    """检测是否为扫描件（前 3 页文字 < 50 字符）"""
    doc = fitz.open(pdf_path)
    total = ""
    for i, page in enumerate(doc):
        if i >= 3:
            break
        total += page.get_text()
    doc.close()
    return len(total.strip()) < 50


def extract_scanned_pdf(pdf_path: str) -> str:
    """扫描件 OCR（需安装 easyocr）"""
    try:
        import easyocr
    except ImportError:
        raise RuntimeError(
            "扫描件 PDF 需要安装 EasyOCR：\n"
            "  uv pip install easyocr"
        )
    import numpy as np
    from PIL import Image

    reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
    doc = fitz.open(pdf_path)
    pages = []

    for page in doc:
        mat = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=mat)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        img_array = np.array(img)
        ocr_result = reader.readtext(img_array)
        page_text = "\n".join(
            [text for (_, text, conf) in ocr_result if conf > 0.4]
        ).strip()
        if page_text:
            pages.append(page_text)

    doc.close()
    text = "\n\n--- 换页 ---\n\n".join(pages)
    logger.info("OCR 提取：%d 字符", len(text))
    return text


# ── 统一入口 ──

def extract_text(pdf_path: str) -> str:
    """PDF 提取入口 — 自动检测类型"""
    if is_scanned_pdf(pdf_path):
        logger.info("检测为扫描件")
        return extract_scanned_pdf(pdf_path)
    logger.info("检测为文字 PDF")
    return extract_pdf(pdf_path)


# ── Word (.docx) 结构化提取 ──

def extract_docx(file_bytes: bytes) -> str:
    """
    从 .docx 提取结构化内容：表格→Markdown表格、标题→#层级、列表→-前缀
    按文档原始顺序输出
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(file_bytes))

    para_by_element = {p._element: p for p in doc.paragraphs}
    table_by_element = {t._element: t for t in doc.tables}

    parts = []
    heading_styles = _get_heading_styles(doc)

    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            para = para_by_element.get(child)
            if para:
                text = _extract_paragraph(para, heading_styles)
                if text:
                    parts.append(text)

        elif tag == 'tbl':
            table = table_by_element.get(child)
            if table:
                md = _table_to_markdown(table)
                if md:
                    parts.append(md)

    result = "\n\n".join(parts)
    logger.info("提取 Word 文档：%d 个元素，%d 字符", len(parts), len(result))
    return result


def _get_heading_styles(doc) -> set:
    from docx.enum.style import WD_STYLE_TYPE
    headings = set()
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.startswith('Heading'):
            headings.add(style.style_id)
    return headings


def _extract_paragraph(para, heading_styles: set) -> str:
    text = para.text.strip()
    if not text:
        return ""

    style_id = para.style.style_id if para.style else ""

    if style_id in heading_styles:
        level = 1
        for char in style_id:
            if char.isdigit():
                level = int(char)
                break
        prefix = "#" * min(level, 6)
        return f"{prefix} {text}"

    if para.style.name == 'List Paragraph' or _has_list_numbering(para):
        return f"- {text}"

    return text


def _has_list_numbering(para) -> bool:
    try:
        numPr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        return numPr is not None
    except Exception:
        return False


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("|" + "|".join([" --- " for _ in range(max_cols)]) + "|")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)
